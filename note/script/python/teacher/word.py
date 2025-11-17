'''
pip install python-docx
'''
from docx import Document
from datetime import datetime
import os
from docx.shared import Pt  # 用于设置字体大小
from docx.oxml.ns import qn
# 获取当前日期（本地时间）
today = datetime.today()
# 格式化为 "年-月-日" 中文格式
formatted_date = today.strftime("%Y年%m月%d日")

file_path = 'E:\A\管控项目资料\日报周报\工程日报017-'+formatted_date+'.docx'

doc = Document('E:\A\管控项目资料\日报周报\日报.docx')

# 今日日报 - 设置内容及格式
today_content = '''
1、完成重点管控人员、管控场所功能的前后端代码编写并打包发布至测试环境。
2、调整并优化风险隐患处置流程的数据模型。
3、完成演示环境系统的安装与部署工作，并进行简单测试。
4、进行AI对话功能与大模型调用测试及适配。
5、调整Dify风险隐患工作流，并完成其调用大模型的测试验证。
6、根据演示环境要求，完成知识库内容的适配调整。
7、新建数据对接库表，抽取12345数据至数据原始库，为后续AI数据分析提供测试数据。
8、完成对大网格系统中所有数据的简单分析统计工作。
'''
# 明日计划 - 设置内容及格式
tomorrow_content = '''
1、开展重点管控人员、管控场所功能的前后端联调与性能优化。
2、继续推进AI对话功能引用大模型的调用测试与深度适配工作。
3、继续接入大网格系统中用户、部门及人口等数据。
4、推进实现情报分析摸排流转功能相关工作。
'''

# 读取所有段落
print("=== 段落内容 ===")
for para in doc.paragraphs:
    print(para.text)

# 读取所有表格
print("\n=== 表格内容 ===")
for table in doc.tables:
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        print(row_data)
        
        
print(doc.paragraphs[1].text)
doc.paragraphs[1].text = '编制时间：'+formatted_date
# 设置字体为宋体，大小为小四（12磅）
for run in doc.paragraphs[1].runs:
    run.font.name = '宋体'
    run.font.size = Pt(12)
    # 解决中文显示问题的额外设置
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
print(doc.paragraphs[1].text)


# 设置今日日报内容
cell_today = doc.tables[0].rows[2].cells[1]
cell_today.text = today_content.strip()  # 先设置文本内容

# 格式化今日日报单元格（宋体小四）
for para in cell_today.paragraphs:
    for run in para.runs:
        run.font.name = '宋体'  # 西文字体
        run.font.size = Pt(12)  # 小四对应12磅
        # 设置中文字体（解决宋体显示问题）
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置明日计划内容
cell_tomorrow = doc.tables[0].rows[4].cells[1]
cell_tomorrow.text = tomorrow_content.strip()  # 先设置文本内容

# 格式化明日计划单元格（宋体小四）
for para in cell_tomorrow.paragraphs:
    for run in para.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        
# 保存修改
if os.path.exists(file_path):
    os.remove(file_path)
doc.save(file_path)











